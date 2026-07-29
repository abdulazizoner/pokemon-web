from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "shopier-sirket-hesabi-ve-operasyon-rehberi.md"
OUTPUT = (
    ROOT
    / "required_fields"
    / "shopier_account.docx"
)

NAVY = "203748"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5F6770"
INK = "1D2329"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
LINE = "CCD5DE"
GOLD = "A66D12"
WHITE = "FFFFFF"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_run_font(
    run,
    *,
    name: str = "Calibri",
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{key}"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:cantSplit")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, width in CELL_MARGIN.items():
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(width))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), LINE)


def set_table_geometry(table, widths: list[int]):
    if sum(widths) != TABLE_WIDTH_DXA:
        raise ValueError(f"Table widths must sum to {TABLE_WIDTH_DXA}: {widths}")

    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        prevent_row_split(row)
        for index, cell in enumerate(row.cells):
            width = widths[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
    set_table_borders(table)


def proportional_widths(rows: list[list[str]]) -> list[int]:
    columns = max(len(row) for row in rows)
    scores = []
    for column in range(columns):
        values = [row[column] if column < len(row) else "" for row in rows]
        score = max(10, min(45, max(len(value) for value in values)))
        scores.append(score)
    total = sum(scores)
    widths = [round(TABLE_WIDTH_DXA * score / total) for score in scores]
    widths[-1] += TABLE_WIDTH_DXA - sum(widths)
    return widths


def add_field(run, instruction: str, fallback: str = ""):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = fallback
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def new_numbering_instance(doc: Document) -> int:
    style = doc.styles["List Number"]
    style_num_id = int(style._element.pPr.numPr.numId.val)
    numbering = doc.part.numbering_part.element
    source_num = next(
        node
        for node in numbering.findall(qn("w:num"))
        if int(node.get(qn("w:numId"))) == style_num_id
    )
    abstract_id = source_num.find(qn("w:abstractNumId")).get(qn("w:val"))
    existing = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    num_id = max(existing) + 1

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), abstract_id)
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    return num_id


def set_paragraph_num_id(paragraph, num_id: int):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = num_pr.find(qn("w:ilvl"))
    if ilvl is None:
        ilvl = OxmlElement("w:ilvl")
        num_pr.append(ilvl)
    ilvl.set(qn("w:val"), "0")
    num = num_pr.find(qn("w:numId"))
    if num is None:
        num = OxmlElement("w:numId")
        num_pr.append(num)
    num.set(qn("w:val"), str(num_id))


def add_hyperlink(paragraph, text: str, url: str):
    relationship = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship)
    run = OxmlElement("w:r")
    run_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_pr.extend([color, underline])
    run.append(run_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE = re.compile(r"(\*\*.+?\*\*|`.+?`|\[[^\]]+\]\(https?://[^)]+\))")


def add_inline(paragraph, text: str):
    cursor = 0
    for match in INLINE.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor : match.start()])
            set_run_font(run, color=INK)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, color=INK, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Consolas", size=9.5, color=DARK_BLUE)
        else:
            link_match = re.match(r"\[([^\]]+)\]\((https?://[^)]+)\)", token)
            if link_match:
                add_hyperlink(paragraph, link_match.group(1), link_match.group(2))
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run, color=INK)


def set_paragraph_outline_none(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    outline = p_pr.find(qn("w:outlineLvl"))
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        p_pr.append(outline)
    outline.set(qn("w:val"), "9")


def add_bottom_border(paragraph, color: str = LINE, size: int = 8):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def setup_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    if "Checklist" not in doc.styles:
        style = doc.styles.add_style("Checklist", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = normal
    checklist = doc.styles["Checklist"]
    checklist.paragraph_format.left_indent = Inches(0.375)
    checklist.paragraph_format.first_line_indent = Inches(-0.188)
    checklist.paragraph_format.space_after = Pt(4)
    checklist.paragraph_format.line_spacing = 1.25

    if "Callout" not in doc.styles:
        style = doc.styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = normal
    callout = doc.styles["Callout"]
    callout.font.color.rgb = rgb(NAVY)
    callout.paragraph_format.space_after = Pt(4)
    callout.paragraph_format.line_spacing = 1.15

    if "Table Text" not in doc.styles:
        style = doc.styles.add_style("Table Text", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = normal
    table_text = doc.styles["Table Text"]
    table_text.font.size = Pt(9.5)
    table_text.paragraph_format.space_after = Pt(2)
    table_text.paragraph_format.line_spacing = 1.1


def setup_sections(doc: Document):
    doc.settings.odd_and_even_pages_header_footer = True

    for section_index, section in enumerate(doc.sections):
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)
        section.different_first_page_header_footer = False

        header_variants = (section.header, section.even_page_header)
        footer_variants = (section.footer, section.even_page_footer)

        for header in header_variants:
            header.is_linked_to_previous = False
            p = header.paragraphs[0]
            p.clear()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            if section_index == 0:
                continue
            left = p.add_run("PROJE DOKÜMANI")
            set_run_font(left, size=8.5, color=NAVY, bold=True)
            right = p.add_run("    |    SHOPIER OPERASYON REHBERİ")
            set_run_font(right, size=8.5, color=MUTED)
            add_bottom_border(p, color=LINE, size=6)

        for footer in footer_variants:
            footer.is_linked_to_previous = False
            fp = footer.paragraphs[0]
            fp.clear()
            fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            fp.paragraph_format.space_before = Pt(0)
            if section_index == 0:
                continue
            label = fp.add_run("Sayfa ")
            set_run_font(label, size=8.5, color=MUTED)
            field_run = fp.add_run()
            set_run_font(field_run, size=8.5, color=MUTED)
            add_field(field_run, "PAGE", "1")


def add_cover(doc: Document):
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(110)
    spacer.paragraph_format.space_after = Pt(0)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    run = kicker.add_run("OPERASYON REHBERİ")
    set_run_font(run, size=10.5, color=GOLD, bold=True)
    run.font.all_caps = True

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("Şirketler İçin Shopier")
    set_run_font(run, size=30, color=NAVY, bold=True)

    title2 = doc.add_paragraph()
    title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title2.paragraph_format.space_after = Pt(10)
    run = title2.add_run("Hesap Açılışı, Ürün Listeleme\nve Operasyon Rehberi")
    set_run_font(run, size=23, color=BLUE, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(30)
    run = subtitle.add_run(
        "Fiziksel koleksiyon kartları için şirket hesabından günlük operasyona"
    )
    set_run_font(run, size=12, color=MUTED, italic=True)

    rule = doc.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rule.paragraph_format.left_indent = Inches(1.35)
    rule.paragraph_format.right_indent = Inches(1.35)
    rule.paragraph_format.space_after = Pt(54)
    add_bottom_border(rule, color=BLUE, size=14)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(4)
    run = meta.add_run("Belge sürümü 1.1  •  30 Temmuz 2026")
    set_run_font(run, size=11, color=NAVY, bold=True)

    scope = doc.add_paragraph()
    scope.alignment = WD_ALIGN_PARAGRAPH.CENTER
    scope.paragraph_format.left_indent = Inches(0.75)
    scope.paragraph_format.right_indent = Inches(0.75)
    scope.paragraph_format.space_after = Pt(16)
    run = scope.add_run(
        "Statik ürün vitrini + Shopier satış akışı için doğrulanmış platform adımları "
        "ve önerilen şirket prosedürleri"
    )
    set_run_font(run, size=9.5, color=MUTED)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.left_indent = Inches(0.8)
    note.paragraph_format.right_indent = Inches(0.8)
    note.paragraph_format.space_before = Pt(36)
    note.paragraph_format.space_after = Pt(0)
    run = note.add_run("ŞİRKET BİLGİLERİ VE HESAP ERİŞİMİ GELMEDEN UYGULANMAMALIDIR")
    set_run_font(run, size=8.5, color=GOLD, bold=True)

    doc.add_section(WD_SECTION.NEW_PAGE)


def add_callout(doc: Document, lines: list[str]):
    for index, line in enumerate(lines):
        paragraph = doc.add_paragraph(style="Callout")
        paragraph.paragraph_format.left_indent = Pt(8)
        paragraph.paragraph_format.right_indent = Pt(8)
        paragraph.paragraph_format.space_before = Pt(8 if index == 0 else 0)
        paragraph.paragraph_format.space_after = Pt(8 if index == len(lines) - 1 else 0)

        p_pr = paragraph._p.get_or_add_pPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), LIGHT_GRAY)
        p_pr.append(shading)

        borders = OxmlElement("w:pBdr")
        edges = ["left", "right"]
        if index == 0:
            edges.append("top")
        if index == len(lines) - 1:
            edges.append("bottom")
        for edge in edges:
            border = OxmlElement(f"w:{edge}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "6")
            border.set(qn("w:space"), "8")
            border.set(qn("w:color"), LINE)
            borders.append(border)
        p_pr.append(borders)
        add_inline(paragraph, line)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(1)


def add_markdown_table(doc: Document, rows: list[list[str]]):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    widths = proportional_widths(rows)
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])

    for row_index, row in enumerate(rows):
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.style = doc.styles["Table Text"]
            add_inline(paragraph, value)
            if row_index == 0:
                set_cell_shading(cell, LIGHT_BLUE)
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = rgb(NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def parse_section(lines: list[str], start_heading: str, end_heading: str | None = None):
    start = next(index for index, line in enumerate(lines) if line == start_heading)
    if end_heading is None:
        return lines[start:]
    end = next(index for index, line in enumerate(lines[start + 1 :], start + 1) if line == end_heading)
    return lines[start:end]


def render_markdown(doc: Document, lines: Iterable[str]):
    lines = list(lines)
    index = 0
    while index < len(lines):
        raw = lines[index].rstrip()
        stripped = raw.strip()
        if not stripped:
            index += 1
            continue

        if stripped.startswith(">"):
            callout = []
            while index < len(lines) and lines[index].strip().startswith(">"):
                text = lines[index].strip()[1:].strip()
                if text:
                    callout.append(text)
                index += 1
            if callout:
                add_callout(doc, callout)
            continue

        if stripped.startswith("|") and index + 1 < len(lines):
            separator = lines[index + 1].strip()
            if separator.startswith("|") and re.search(r"\|\s*:?-{3,}", separator):
                table_rows = []
                header = [cell.strip() for cell in stripped.strip("|").split("|")]
                table_rows.append(header)
                index += 2
                while index < len(lines) and lines[index].strip().startswith("|"):
                    table_rows.append(
                        [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
                    )
                    index += 1
                if all(len(row) == len(header) for row in table_rows):
                    add_markdown_table(doc, table_rows)
                continue

        heading = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading:
            level = len(heading.group(1))
            paragraph = doc.add_paragraph(style=f"Heading {level}")
            add_inline(paragraph, heading.group(2))
            index += 1
            continue

        checklist = re.match(r"^-\s+\[\s*\]\s+(.+)$", stripped)
        if checklist:
            paragraph = doc.add_paragraph(style="Checklist")
            run = paragraph.add_run("☐ ")
            set_run_font(run, color=BLUE, bold=True)
            add_inline(paragraph, checklist.group(1))
            index += 1
            continue

        bullet = re.match(r"^-\s+(.+)$", stripped)
        if bullet:
            paragraph = doc.add_paragraph(style="List Bullet")
            add_inline(paragraph, bullet.group(1))
            index += 1
            continue

        numbered = re.match(r"^\d+\.\s+(.+)$", stripped)
        if numbered:
            items = []
            while index < len(lines):
                item = re.match(r"^\d+\.\s+(.+)$", lines[index].strip())
                if not item:
                    break
                items.append(item.group(1))
                index += 1
            num_id = new_numbering_instance(doc)
            for item in items:
                paragraph = doc.add_paragraph(style="List Number")
                set_paragraph_num_id(paragraph, num_id)
                add_inline(paragraph, item)
            continue

        if stripped.startswith("`") and stripped.endswith("`") and len(stripped) > 2:
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.25)
            run = paragraph.add_run(stripped[1:-1])
            set_run_font(run, name="Consolas", size=9.5, color=DARK_BLUE)
            index += 1
            continue

        paragraph_lines = [stripped]
        index += 1
        while index < len(lines):
            candidate = lines[index].strip()
            if (
                not candidate
                or candidate.startswith("#")
                or candidate.startswith(">")
                or candidate.startswith("|")
                or re.match(r"^-\s+", candidate)
                or re.match(r"^\d+\.\s+", candidate)
            ):
                break
            paragraph_lines.append(candidate)
            index += 1
        paragraph = doc.add_paragraph()
        add_inline(paragraph, " ".join(paragraph_lines))


def add_intro_and_toc(doc: Document, lines: list[str]):
    intro = parse_section(lines, "## 1. Belgenin amacı ve kapsamı", "## 2. Son doğrulama tarihi ve değişiklik yönetimi")
    verified = parse_section(lines, "## 2. Son doğrulama tarihi ve değişiklik yönetimi", "## 3. İçindekiler")
    render_markdown(doc, intro)
    render_markdown(doc, verified)

    doc.add_page_break()
    toc_heading = doc.add_paragraph()
    toc_heading.paragraph_format.space_before = Pt(0)
    toc_heading.paragraph_format.space_after = Pt(12)
    set_paragraph_outline_none(toc_heading)
    run = toc_heading.add_run("3. İçindekiler")
    set_run_font(run, size=18, color=BLUE, bold=True)
    toc_pages = {
        1: 2, 2: 2, 3: 3, 4: 4, 5: 4, 6: 5, 7: 7, 8: 7, 9: 8, 10: 9,
        11: 10, 12: 10, 13: 11, 14: 12, 15: 12, 16: 13, 17: 14, 18: 15,
        19: 15, 20: 16, 21: 17, 22: 17, 23: 18, 24: 19, 25: 20, 26: 20,
        27: 21, 28: 22, 29: 23, 30: 23, 31: 24, 32: 24, 33: 25, 34: 26,
        35: 27, 36: 28, 37: 28, 38: 29,
    }
    entries = []
    for line in lines:
        match = re.match(r"^##\s+(\d+)\.\s+(.+)$", line)
        if match:
            entries.append((int(match.group(1)), match.group(2)))
    for section_number, label in entries:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.keep_together = True
        paragraph.paragraph_format.tab_stops.add_tab_stop(
            Inches(6.25), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
        )
        run = paragraph.add_run(f"{section_number}. {label}")
        set_run_font(run, size=8.5, color=INK)
        page_run = paragraph.add_run(f"\t{toc_pages[section_number]}")
        set_run_font(page_run, size=8.5, color=MUTED)
    doc.add_page_break()


def set_document_properties(doc: Document):
    props = doc.core_properties
    props.title = "Şirketler İçin Shopier Hesabı Açılışı, Ürün Listeleme ve Operasyon Rehberi"
    props.subject = "Shopier ticari hesap, ürün, kargo, sipariş ve teslim operasyon rehberi"
    props.author = "Proje ekibi"
    props.keywords = "Shopier, şirket hesabı, koleksiyon kartları, ürün listeleme, kargo, iade"
    props.comments = "Son doğrulama: 30 Temmuz 2026. Şirket bilgileri geçici olarak beklenmektedir."

    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def audit_document(doc: Document):
    assert len(doc.sections) == 2
    for section in doc.sections:
        assert round(section.page_width.inches, 3) == 8.5
        assert round(section.page_height.inches, 3) == 11
        for margin in (
            section.top_margin,
            section.right_margin,
            section.bottom_margin,
            section.left_margin,
        ):
            assert round(margin.inches, 3) == 1.0
        assert round(section.header_distance.inches, 3) == 0.492
        assert round(section.footer_distance.inches, 3) == 0.492

    assert doc.sections[0].header.paragraphs[0].text == ""
    assert doc.sections[0].footer.paragraphs[0].text == ""
    assert "SHOPIER OPERASYON REHBERİ" in doc.sections[1].header.paragraphs[0].text
    assert doc.sections[1].footer.paragraphs[0].text.startswith("Sayfa ")

    normal = doc.styles["Normal"]
    assert normal.font.name == "Calibri"
    assert normal.font.size.pt == 11
    assert normal.paragraph_format.space_after.pt == 6

    for table in doc.tables:
        tbl_w = table._tbl.tblPr.find(qn("w:tblW"))
        tbl_ind = table._tbl.tblPr.find(qn("w:tblInd"))
        assert tbl_w is not None and int(tbl_w.get(qn("w:w"))) == TABLE_WIDTH_DXA
        assert tbl_ind is not None and int(tbl_ind.get(qn("w:w"))) == TABLE_INDENT_DXA
        grid_total = sum(int(col.get(qn("w:w"))) for col in table._tbl.tblGrid)
        assert grid_total == TABLE_WIDTH_DXA


def main():
    markdown = SOURCE.read_text(encoding="utf-8")
    lines = markdown.splitlines()
    body_start = next(index for index, line in enumerate(lines) if line.startswith("## 4. "))

    doc = Document()
    setup_styles(doc)
    set_document_properties(doc)
    add_cover(doc)
    add_intro_and_toc(doc, lines)
    render_markdown(doc, lines[body_start:])
    setup_sections(doc)
    audit_document(doc)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")


if __name__ == "__main__":
    main()
